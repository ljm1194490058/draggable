from django.shortcuts import render, redirect
import requests
from bs4 import BeautifulSoup
import base64
from django.http import JsonResponse
from io import BytesIO
from docx import Document
import kaggle
import os
from .forms import MyModelForm
from django.conf import settings
from django.http import HttpResponse
import pandas as pd
import numpy as np
import csv
from sklearn.preprocessing import LabelEncoder, OneHotEncoder, MinMaxScaler, StandardScaler
from apyori import apriori
from mlxtend.preprocessing import TransactionEncoder
from mlxtend.frequent_patterns import apriori, association_rules
from django.contrib import messages
from sklearn.linear_model import LinearRegression
import seaborn as sns
import matplotlib.pyplot as plt
from django.views.decorators.csrf import csrf_exempt
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.keys import Keys


# 当前选中的数据集
select_file = ''
# 当前选中数据集的dataframe
select_df = pd.DataFrame()
# 原本数据的dataframe
index_df = pd.DataFrame()
# 预处理后的df
processing_df = pd.DataFrame()


# 数据预处理模块  post是上传文件
def child(request):
    global select_file
    global select_df
    if request.method == 'POST':
        file = request.FILES['file']
        # 更新选中的数据集
        select_file = str(file)
        ##########区分csv文件和excel文件
        # 获取文件扩展名
        file_extension = select_file.split('.')[-1]
        # 如果是CSV文件，则使用read_csv函数读取数据
        if file_extension == 'csv':
            df = pd.read_csv(file)
        elif file_extension == 'txt':
            df = pd.read_csv(file, delimiter='\t')

        else:
            # 读取上传的数据
            df = pd.read_excel(file)
        data = df
        select_df = df
        # 获取数据集的形状
        shape = df.shape
        # 输出数据集的形状和大小
        data_shape = str(shape[0]) + "行x" + str(shape[1]) + "列"
        print("预处理模块的读取文件成功！")
        filename = select_file
        return render(request, "child.html", {'data':data, 'file_title':filename, 'data_shape':data_shape})

    else:
        filename = select_file
        data = select_df
        return render(request, 'child.html', {'data': data, "filename":filename})


###################针对table.html##########################
# 文本文件转csv文件
def txt_to_csv(txt_file_path, csv_file_path=None):
    # 获取TXT文件路径和名称
    if not csv_file_path:
        txt_file_name = str(txt_file_path)
        csv_file_name = txt_file_name.split('.')[0] + '.csv'

        # Read the TXT file data into a Pandas DataFrame
    with open(txt_file_path, 'r') as f:
        lines = f.readlines()
    data = [line.strip().split('\t') for line in lines]
    df = pd.DataFrame(data)

    return df

# 上传文件
def upload_file(request):
    global select_file
    global select_df
    global index_df
    if request.method == 'POST':
        file = request.FILES['file']
        # 更新选中的数据集
        select_file = str(file)
        ##########区分csv文件和excel文件
        # 获取文件扩展名
        file_extension = select_file.split('.')[-1]

        # 如果是CSV文件，则使用read_csv函数读取数据
        if file_extension == 'csv':
            df = pd.read_csv(file)
        elif file_extension == 'txt':
            df = pd.read_csv(file, delimiter='\t')
        else:
        # 读取上传的数据
            df = pd.read_excel(file)
        print("读取文件成功！")
        # 更新选中数据集的dataframe
        select_df = df
        index_df = df
        # 动态获取数据集的键和值
        columns = df.columns.tolist()
        # 只取几列进行显示
        values = df.values.tolist()

        # 基本统计信息
        # 获取数据集的形状
        shape = df.shape
        # 输出数据集的形状和大小
        data_shape = str(shape[0]) + "行x" + str(shape[1]) + "列"
        # 获取每列的数据类型
        data_types = df.dtypes.tolist()
        # 获取每一列的缺失值数量
        missing_values = df.isnull().sum().tolist()
        # 计算每一列的离群值
        num_outliers = []
        for col in df.columns:
            if df[col].dtype == 'float' or df[col].dtype == 'int' or df[col].dtype == 'int64':
                # 使用 Z-分数方法计算离群值， 如果某个值的 Z-分数大于 3，则将其视为异常值
                z_scores = np.abs((df[col] - df[col].mean()) / df[col].std())
                num_outliers.append(sum(z_scores > 3))
            else:
                num_outliers.append('该数据类型不支持该处理！')

        # 基本数值统计量
        avg_list = []
        for feature in df.columns:
            if df[feature].dtype == 'float' or df[feature].dtype == 'int' or df[feature].dtype == 'int64':
                median = round(df[feature].mean(), 2)
            else:
                median = '该数据类型不支持该处理！'
            avg_list.append(median)
        # 标准差
        std_list = []
        for feature in df.columns:
            if df[feature].dtype == 'float' or df[feature].dtype == 'int' or df[feature].dtype == 'int64':
                median = round(df[feature].std(), 2)
            else:
                median = '该数据类型不支持该处理！'
            std_list.append(median)
        min_list = []
        for feature in df.columns:
            if df[feature].dtype == 'float' or df[feature].dtype == 'int' or df[feature].dtype == 'int64':
                median = round(df[feature].min(), 2)
            else:
                median = '该数据类型不支持该处理！'
            min_list.append(median)
        max_list = []
        for feature in df.columns:
            if df[feature].dtype == 'float' or df[feature].dtype == 'int' or df[feature].dtype == 'int64':
                median = round(df[feature].max(), 2)
            else:
                median = '该数据类型不支持该处理！'
            max_list.append(median)
        median_list = []
        for feature in df.columns:
            if df[feature].dtype == 'float' or df[feature].dtype == 'int' or df[feature].dtype == 'int64':
                median = round(df[feature].median(), 2)
            else:
                median = '该数据类型不支持该处理！'
            median_list.append(median)


        filename = file.name
        filepath = os.path.join(settings.MEDIA_ROOT, filename)
        with open(filepath, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)
        file_url = '/media/' + filename
        print("fi:", file_url)

        # 将数据传递给模板
        context = {
            'columns': columns,
            'values': values,
            'file_url': file_url,
            'file_title':filename,
            'avg_list':avg_list,
            'std_list':std_list,
            'max_list':max_list,
            'min_list':min_list,
            'median_list':median_list,
            'data_shape':data_shape,
            'data_types':data_types,
            'missing_values':missing_values,
            'num_outliers':num_outliers
        }

        return render(request, 'table.html', context)
    else:
        if select_file != '':
            # 动态获取数据集的键和值
            columns = index_df.columns.tolist()

            df = index_df
            # 只取几列进行显示
            values = df.values.tolist()

            # 基本统计信息
            # 获取数据集的形状
            shape = df.shape
            # 输出数据集的形状和大小
            data_shape = str(shape[0]) + "行x" + str(shape[1]) + "列"
            # 获取每列的数据类型
            data_types = df.dtypes.tolist()
            # 获取每一列的缺失值数量
            missing_values = df.isnull().sum().tolist()
            # 计算每一列的离群值
            num_outliers = []
            for col in df.columns:
                if df[col].dtype == 'float' or df[col].dtype == 'int' or df[col].dtype == 'int64':
                    # 使用 Z-分数方法计算离群值， 如果某个值的 Z-分数大于 3，则将其视为异常值
                    z_scores = np.abs((df[col] - df[col].mean()) / df[col].std())
                    num_outliers.append(sum(z_scores > 3))
                else:
                    num_outliers.append('该数据类型不支持该处理！')

            # 基本数值统计量
            avg_list = []
            for feature in df.columns:
                if df[feature].dtype == 'float' or df[feature].dtype == 'int'or df[feature].dtype == 'int64':
                    median = round(df[feature].mean(), 2)
                else:
                    median = '该数据类型不支持该处理！'
                avg_list.append(median)
            # 标准差
            std_list = []
            for feature in df.columns:
                if df[feature].dtype == 'float' or df[feature].dtype == 'int' or df[feature].dtype == 'int64':
                    median = round(df[feature].std(), 2)
                else:
                    median = '该数据类型不支持该处理！'
                std_list.append(median)
            min_list = []
            for feature in df.columns:
                if df[feature].dtype == 'float' or df[feature].dtype == 'int' or df[feature].dtype == 'int64':
                    median = round(df[feature].min(), 2)
                else:
                    median = '该数据类型不支持该处理！'
                min_list.append(median)
            max_list = []
            for feature in df.columns:
                if df[feature].dtype == 'float' or df[feature].dtype == 'int' or df[feature].dtype == 'int64':
                    median = round(df[feature].max(), 2)
                else:
                    median = '该数据类型不支持该处理！'
                max_list.append(median)
            median_list = []
            for feature in df.columns:
                if df[feature].dtype == 'float' or df[feature].dtype == 'int' or df[feature].dtype == 'int64':
                    median = round(df[feature].median(), 2)
                else:
                    median = '该数据类型不支持该处理！'
                median_list.append(median)

            file = select_file
            # filepath = os.path.join(settings.MEDIA_ROOT, file)
            # with open(filepath, 'wb+') as destination:
            #     for chunk in file.chunks():
            #         destination.write(chunk)
            file_url = '/media/' + file
            print("fi:", file_url)

            # 将数据传递给模板
            context = {
                'columns': columns,
                'values': values,
                'file_url': file_url,
                'file_title': file,
                'avg_list': avg_list,
                'std_list': std_list,
                'max_list': max_list,
                'min_list': min_list,
                'median_list': median_list,
                'data_shape': data_shape,
                'data_types': data_types,
                'missing_values': missing_values,
                'num_outliers': num_outliers
            }
            return render(request, 'table.html', context)
        else:
            print("eror!!!!")
        return render(request, 'table.html')


# 下载文件
def download_file(request, filename):
    filepath = os.path.join(settings.MEDIA_ROOT, filename)
    with open(filepath, 'rb') as file:
        response = HttpResponse(file.read(), content_type='application/force-download')
        response['Content-Disposition'] = 'attachment; filename=%s' % filename
        response['X-Sendfile'] = filepath
        return response


###################针对dashboard.html######################
def line_chart(request):
    try:
        print('select_file',select_file)
        df = index_df
        # 动态获取数据集的键和值

        # 提取每个int特征为列表
        all_int_list = []
        all_string_list = []
        columns = []
        all_column = []

        for i in range(0, df.shape[0]):
            all_column.append(i+1)
        for i, col in enumerate(df.columns):
            if df[col].dtype == 'float' or df[col].dtype == 'int':
                list1 = df.iloc[:, i].tolist()
                # 每条折线的数据
                all_int_list.append(list1)
                # 有哪些折线
                columns.append(col)
            else:
                pie_data_index = list(df[col].value_counts().index)
                pie_data = list(df[col].value_counts())
                for i in range(len(pie_data)):
                    dic = {}
                    dic['name'] = pie_data_index[i]
                    dic['value'] = pie_data[i]
                    all_string_list.append(dic)

    #  散点图
        san_list = []
        for j, k in zip(df[columns[0]], df[columns[1]]):
            san_list.append([j, k])

        context = {
            'all_int_list':all_int_list,
            'all_column':all_column,
            'columns':columns,
            'san_list':san_list
        }
        return render(request, 'dashboard.html', context)
    except:
        return render(request, 'dashboard.html')




# Create your views here.

research_all_df = pd.DataFrame()
def research(request):
    global research_all_df
    if 'research_all' in request.POST:
        url = "https://archive.ics.uci.edu/ml/datasets.php"
        research_df = pd.read_html(url, header=0)[5]
        research_df = research_df.rename(columns={'Name': 'name'})

        # 发送 HTTP 请求
        url = 'https://archive.ics.uci.edu/ml/datasets.php'
        response = requests.get(url)


        # 解析 HTML
        soup = BeautifulSoup(response.text, 'html.parser')
        df1 = pd.DataFrame(columns=['name', 'href'])
        # 提取数据集链接
        datasets = soup.find_all('table', cellpadding='5')[0].find_all('a')
        for dataset in datasets:
            name = dataset.text.strip()
            link = 'https://archive.ics.uci.edu/ml/' + dataset['href']
            # print(dataset.text, link)
            if name != '':
                df1 = df1.append({'name': name, 'href': link}, ignore_index=True)
        # pd.set_option('display.max_rows', None)
        # 设置display.max_columns属性为None，以展示所有列
        # pd.set_option('display.max_columns', None)
        #
        # # 将最大列宽设置为-1，以便完全显示每列
        # pd.set_option('display.max_colwidth', -1)
        df1 = df1.iloc[7:, ]

        merged_df = pd.merge(research_df, df1, on='name')
        print("merged_df", merged_df.head())
        research_all_df = merged_df
        research_column = merged_df.columns.tolist()
        research_value = merged_df.values.tolist()
        return render(request, "research.html", {"df_value":research_value, "df_column":research_column})

    # 查找单个数据
    elif 'single' in request.POST:
        name1 = request.POST.get('research_input')
        #对字符串进行操作
        name1 = name1.replace('\t', '')
        name1 = name1.replace('\n', '')

        row_data = research_all_df.loc[research_all_df['name'] == name1]
        #############[0]就是个神来之笔#################
        row_data = row_data.values.tolist()[0]
        print("row",row_data)
        research_column = research_all_df.columns.tolist()
        research_value = research_all_df.values.tolist()
        return render(request, "research.html", {"df_value":research_value, "df_column":research_column, 'row_data':row_data})
    else:
        research_column = research_all_df.columns.tolist()
        research_value = research_all_df.values.tolist()
        return render(request, "research.html", {"df_value":research_value, "df_column":research_column})





def index(request):
    return render(request, "research.html")


def sandian(request):
    if request.method == 'POST':
        print('select_file', select_file)
        df = index_df
        # 动态获取数据集的键和值

        # 提取每个int特征为列表
        all_int_list = []
        all_string_list = []
        columns = []
        all_column = []

        for i in range(0, df.shape[0]):
            all_column.append(i + 1)
        for i, col in enumerate(df.columns):
            if df[col].dtype == 'float' or df[col].dtype == 'int':
                list1 = df.iloc[:, i].tolist()
                # 每条折线的数据
                all_int_list.append(list1)
                # 有哪些折线
                columns.append(col)
            else:
                pie_data_index = list(df[col].value_counts().index)
                pie_data = list(df[col].value_counts())
                for i in range(len(pie_data)):
                    dic = {}
                    dic['name'] = pie_data_index[i]
                    dic['value'] = pie_data[i]
                    all_string_list.append(dic)

        #  散点图
        selected_columns = request.POST.getlist('item')
        san_list = []
        for j, k in zip(index_df[selected_columns[0]], index_df[selected_columns[1]]):
            san_list.append([j, k])

        context = {
            'all_int_list': all_int_list,
            'all_column': all_column,
            'columns': columns,
            'san_list': san_list
        }
        return render(request, 'dashboard.html', context)


def upload_process(request):
    global processing_df
    if request.method == 'POST':
        data = request.POST.dict()

        # 处理表单数据
        name1_value = data.get('name1')
        name2_value = data.get('name2')
        name3_value = data.get('name3')
        name4_value = data.get('name4')
        df = select_df
        for col in df.columns:
            if df[col].dtype == 'int' or df[col].dtype == 'float':
                df[col] = df[col].fillna(df[col].mean())
            else:
                df[col] = df[col].fillna('NANA')
        # df.fillna('NAN', inplace=True)  # 将缺失值替换为-1
        if name1_value == '异常值处理':
            # 遍历每个数值列
            for col in df.select_dtypes(include=['number']).columns:
                # 计算该列的均值和标准差
                mean = df[col].mean()
                std = df[col].std()

                # 找出该列中所有绝对值大于3倍标准差的值
                outliers = df[(df[col] - mean).abs() > 3 * std][col]

                # 用该列的中位数来替换这些异常值
                if not outliers.empty:
                    median = df[col].median()
                    df.loc[outliers.index, col] = median

        if name3_value == '数值化处理':
            # 对每个object和category列进行one-hot编码
            object_cols = list(df.select_dtypes(include=['object', 'category']).columns)

            # 对每个object和category列进行one-hot编码
            for col in object_cols:
                dummies = pd.get_dummies(df[col], prefix=col)
                df = pd.concat([df, dummies], axis=1)
                df.drop(col, axis=1, inplace=True)

        if name4_value == '归一化处理':
            mms = MinMaxScaler()
            for col in df.select_dtypes(include=['number']).columns:
                df[col] = mms.fit_transform(df[[col]])

        try:
            df.drop('Id', axis=1, inplace=True)
        except KeyError:
            pass
        try:
            df.drop('Id', axis=1, inplace=True)
        except KeyError:
            pass
        # 在此处执行对表单数据的处理逻辑

        data = select_df
        processing_df = df
        context = {
            "yuchuli_df":df,
            'data': data,
            'file_title': select_file,
        }
        # 返回一个 JSON 响应
        return render(request, 'child.html', context)

guanlian_column = []
guanlian_value = []
huigui_xiangguan = []
huigui_imd = ''
g_huigui_x = ''
g_huigui_y = ''

def data_analytics(request):
    global guanlian_column
    global guanlian_value
    global huigui_imd
    global g_huigui_x
    global g_huigui_y
    global huigui_xiangguan
    if 'guanlian' in request.POST:

        # 先把所有的特征名称保留好
        df = processing_df
        columns = []
        for i, col in enumerate(df.columns):
            if df[col].dtype == 'float' or df[col].dtype == 'int' or df[col].dtype == 'int64':
                columns.append(col)
        # 这里才开始 a
        selected_columns = request.POST.getlist('item')

        df = df[selected_columns]
        print("selected_columns——df:", df)
        try:
            rules = assoc_rules(df, min_supp=0.1, min_conf=0.5)
        except Exception as e:
            error_message = str("数据集过小或不符合关联规则使用要求!")
            js_code = f"<script>alert('{error_message}'); window.location.href='/data_analytics';</script>"
            return HttpResponse(js_code)

        rules_df = pd.DataFrame(rules)
        rule_column = rules_df.columns.tolist()
        rule_df = rules_df.values.tolist()
        print("ru_df",rule_df)
        guanlian_column = rule_column
        guanlian_value = rule_df
        if rules_df.empty:
            error_message = str("所选特征不足或不符合关联规则使用要求!")
            js_code = f"<script>alert('{error_message}'); window.location.href='/data_analytics';</script>"
            return HttpResponse(js_code)
        else:
            context = {
                "rule_column":rule_column,
                "rule_df":rule_df,
                'columns': columns,
                'xiangguan':huigui_xiangguan,
                "image_path": huigui_imd,
                "huigui_x":g_huigui_x,
                "huigui_y":g_huigui_y,
            }
            return render(request, 'notifications.html', context)


    elif 'huigui' in request.POST:
        df = processing_df
        columns = []
        for i, col in enumerate(df.columns):
            if df[col].dtype == 'float' or df[col].dtype == 'int' or df[col].dtype == 'int64':
                columns.append(col)

        selected_columns = request.POST.getlist('item')
        df = df[selected_columns]
        print("selected_columns——df:", df)
        huigui_x = selected_columns[0]
        g_huigui_x = huigui_x
        if df.shape[1] == 1:
            error_message = str("所选特征不足!")
            js_code = f"<script>alert('{error_message}'); window.location.href='/data_analytics';</script>"
            return HttpResponse(js_code)


        else:
            huigui_y = selected_columns[1]
            g_huigui_y = huigui_y
                # 创建线性回归模型
            lr = LinearRegression()

            # 拟合模型
            X = df.iloc[:, 0].values.reshape(-1, 1)
            y = df.iloc[:, 1].values
            lr.fit(X, y)
            # 创建 Figure 对象并设置大小
            plt.clf()
            plt.rcParams['font.sans-serif'] = ['SimHei']
            fig, ax = plt.subplots(figsize=(5, 4))
            # 绘制回归曲线和散点图
            sns.regplot(x=X, y=y, data=df)
            plt.xlabel('X轴特征')
            plt.ylabel('Y轴特征')
            plt.title('线性回归图像')
            xiangguan = lr.coef_.tolist()
            huigui_xiangguan = xiangguan
            print("相关系数：", lr.coef_)

            # 图片生成
            buffer = BytesIO()
            plt.savefig(buffer)
            plot_data = buffer.getvalue()
            imb = base64.b64encode(plot_data)
            ims = imb.decode()
            imd = "data:image/png;base64," + ims
            # 显示图形和相关系数
            huigui_imd = imd

            context = {
                'columns': columns,
                'xiangguan':xiangguan,
                "image_path": imd,
                "huigui_x":huigui_x,
                "huigui_y":huigui_y,
                "rule_column": guanlian_column,
                "rule_df": guanlian_value,

            }
            return render(request, 'notifications.html', context)

    elif 'xiangguan' in request.POST:
        df = processing_df
        columns = []
        for i, col in enumerate(df.columns):
            if df[col].dtype == 'float' or df[col].dtype == 'int' or df[col].dtype == 'int64':
                columns.append(col)

        selected_columns = request.POST.getlist('item')
        df = df[selected_columns]

        plt1 = correlation_analysis(df)
        # 图片生成
        buffer = BytesIO()
        plt1.savefig(buffer)
        plot_data = buffer.getvalue()
        imb = base64.b64encode(plot_data)
        ims = imb.decode()
        imd = "data:image/png;base64," + ims
        # 显示图形和相关系数
        context = {
            'columns': columns,
            "xiangguan_path": imd,
            "rule_df": guanlian_value,
            "rule_column": guanlian_column,
            'xiangguan': huigui_xiangguan,
            "image_path": huigui_imd,
            "huigui_x": g_huigui_x,
            "huigui_y": g_huigui_y,
        }
        return render(request, 'notifications.html', context)
    else:
        print('select_file', select_file)
        df = processing_df

        columns = []
        for i, col in enumerate(df.columns):
            if df[col].dtype == 'float' or df[col].dtype == 'int' or df[col].dtype == 'int64':
                columns.append(col)

        context = {
            'columns': columns,
        }
        return render(request, 'notifications.html', context)




# 关联规则
def assoc_rules(df, min_supp, min_conf):
    # Convert data to list of lists
    transactions = []
    for i in range(len(df)):
        transactions.append([str(df.values[i, j]) for j in range(len(df.columns))])

    # Run Apriori algorithm
    te = TransactionEncoder()
    te_ary = te.fit(transactions).transform(transactions)
    df = pd.DataFrame(te_ary, columns=te.columns_)

    frequent_itemsets = apriori(df, min_support=min_supp, use_colnames=True)
    rules = association_rules(frequent_itemsets, metric="confidence", min_threshold=min_conf)

    # Return results
    return rules

# 相关性分析
def correlation_analysis(data):
    '''
    对给定数据集的所有数值特征进行相关性分析

    参数：
    data：DataFrame 类型，原始数据集

    返回值：
    None
    '''

    # 计算相关系数矩阵
    corr_matrix = data.corr()

    plt.clf()
    plt.rcParams['font.sans-serif'] = ['SimHei']

    # 绘制相关系数矩阵热力图
    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm')
    plt.title('Correlation Matrix of Numeric Features')
    return plt


def export_to_word(request):
    # 获取要导出的 HTML 内容
    content = request.POST.get('content')

    # 创建 Word 文档
    document = Document()
    section = document.sections[0]
    section.orientation = 1 # 设置页面方向为纵向
    section.top_margin = 914400 # 设置页面上边距
    section.bottom_margin = 914400 # 设置页面下边距
    section.left_margin = 1143000 # 设置页面左边距
    section.right_margin = 1143000 # 设置页面右边距


    chart1_exists = False
    chart2_exists = False
    # 解析 HTML 并将其转换为 Word 文档元素
    soup = BeautifulSoup(content, 'html.parser')
    for tag in soup.find_all():
        if tag.name in ('h1', 'h2', 'h3'):
            # 添加标题
            document.add_heading(tag.text, level=int(tag.name[-1]))
        elif tag.name == 'table':
            # 添加表格
            rows = tag.find_all('tr')
            cols = len(rows[0].find_all('td'))
            if cols == 0: # 当表格为空时，忽略错误
                continue
            table = document.add_table(rows=len(rows), cols=cols)
            for i, row in enumerate(rows):
                cells = row.find_all('td')
                for j, cell in enumerate(cells):
                    table.cell(i, j).text = cell.text.strip()
        elif tag.name == 'input':
            # 忽略 input 标签
            continue

        elif tag.name == 'div':
            if tag.get('style') is not None and 'display:none' in tag['style']:
                continue
            if tag.get('id') == 'chart':
                img_data = base64.b64decode(tag.img['src'].split(',')[1])
                img_buffer = BytesIO(img_data)
                document.add_picture(img_buffer)
            if tag.get('id') == 'chart1':
                try:
                    img_data = base64.b64decode(tag.img['src'].split(',')[1])
                    img_buffer = BytesIO(img_data)
                    document.add_picture(img_buffer)
                except:
                    continue
        else:
            # 判断文本段落是否可见
            if tag.get('style') is not None and 'display:none' in tag['style']:
                continue
            # 添加文本段落
            document.add_paragraph(tag.text)

    # 将文档保存为二进制流
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # 返回响应，并设置文件名和文件类型
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="data_analysis_report.docx"'
    return response



# 关联规则输出属性意思：
# antecedents: 规则的前提条件，即规则左侧的项集。
# consequents: 规则的结论，即规则右侧的项集。
# antecedent support: 前提条件的支持度。
# consequent support: 结论的支持度。
# support: 规则的支持度，即同时出现在整个数据集中的项集的比例。
# confidence: 规则的置信度，即在前提条件出现的情况下，结论出现的概率。
# lift: 规则的提升度，表示前提条件的出现会对结论出现的影响程度，如果lift大于1，则表示前提条件出现提高了结论的出现概率，否则减小了结论的出现概率。
# leverage: 规则的杠杆值，衡量前提条件和结论出现的相关性。
# conviction: 规则的确信度，衡量前提条件出现与结论未出现之间的相关性。

